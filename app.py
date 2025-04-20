from flask import Flask, render_template, request, send_file
from docx import Document
import os
import sys
from datetime import datetime
from io import BytesIO

app = Flask(__name__)

# Функция для получения правильного пути к файлу при запуске через PyInstaller или сервер
def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Функция для преобразования даты в формат "«день» месяц год г."
def format_date(date_obj):
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    day = date_obj.day
    month = months[date_obj.month]
    year = date_obj.year
    return f"«{day}» {month} {year} г."

# Функция для преобразования даты в формат "дд.мм.гггг г."
def format_tour_date(date_str):
    date_obj = datetime.strptime(date_str, "%d.%m.%Y")
    day = date_obj.day
    month = date_obj.month  
    year = date_obj.year
    return f"{day:02d}.{month:02d}.{year} г."

# Функция для замены плейсхолдеров в тексте
def replace_placeholders(text, replacements):
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text

# Функция для определения правильного склонения слова "ночь"
def format_nights(night_num):
    num = int(night_num)
    last_two_digits = num % 100
    last_digit = num % 10

    if 10 <= last_two_digits <= 20:
        return f"{num} ночей"
    elif last_digit == 1:
        return f"{num} ночь"
    elif 2 <= last_digit <= 4:
        return f"{num} ночи"
    else:
        return f"{num} ночей"

# Текст про лечение для класса "Люкс"
luxe_treatment_text = """Лечение, включённое в номер "Люкс" на человека:
• Консультация врача (в начале курса и при необходимости)
• Пантогематоген — 50 мл натощак (ежедневно утром, на тощак)
• Две пантовые ванны с гидромассажем в день (ежедневно) утром и вечером
• Фитобочка с алтайскими травами (ежедневно) 1 раз
• Фито-чай с мёдом (ежедневно) после фитобочки
• Маска с пантовым отваром (ежедневно) 1 раз (после ванны)
• Пантовые обёртывания в 5 дней 1 раз"""

# Текст про лечение для класса "Стандарт"
standard_treatment_text = """Лечение, включённое в номер "Стандарт" на человека:
• Консультация врача (в начале курса и при необходимости)
• Пантогематоген — 50 мл натощак (ежедневно утром, на тощак)
• Одна пантовая ванна с гидромассажем (ежедневно)
• Фитобочка с алтайскими травами (ежедневно 1 раз)
• Фито-чай с мёдом (ежедневно, после фитобочки)"""

@app.route("/", methods=["GET", "POST"])
def generate_doc():
    if request.method == "POST":
        full_name = request.form["full_name"]
        contract_number = request.form["contract_number"]
        id_num = request.form["id_num"]
        night_num = request.form["night_num"]
        start_date = request.form["start_date"]
        end_date = request.form["end_date"]
        date_birth = request.form["date_birth"]
        iin = request.form["iin"]
        with_treatment = request.form.get("with_treatment") == "on"

        # Формируем список туристов, начиная с full_name
        tourists = [f"1) {full_name}"]
        i = 0
        while True:
            tourist_name = request.form.get(f"tourist_{i}", "").strip()
            if not tourist_name:
                break
            tourists.append(f"{i + 2}) {tourist_name}")
            i += 1
        tourists_list = "\n".join(tourists)

        # Формируем список номеров
        rooms = []
        room_classes = set()  # Множество для отслеживания классов номеров
        i = 0
        while True:
            capacity = request.form.get(f"room_capacity_{i}")
            room_class = request.form.get(f"room_class_{i}")
            quantity = request.form.get(f"room_quantity_{i}")
            if not capacity or not room_class or not quantity:
                break
            room_text = f"{capacity}-местный «{room_class}» *{quantity}шт номера"
            rooms.append(room_text)
            room_classes.add(room_class)  # Добавляем класс номера в множество
            i += 1
        rooms_list = "\n\n".join(rooms) if rooms else "Не указано"

        # Формируем текст про лечение
        treatment_details = []
        if "Стандарт" in room_classes:
            treatment_details.append(standard_treatment_text)
        if "Люкс" in room_classes:
            treatment_details.append(luxe_treatment_text)
        treatment_details_text = "\n\n".join(treatment_details) if treatment_details else "Лечение не включено"

        try:
            # Получаем текущую дату
            current_date = datetime.today()
            formatted_date = format_date(current_date)

            # Преобразуем даты
            formatted_start_date = format_tour_date(start_date)
            formatted_end_date = format_tour_date(end_date)
            formatted_birth_date = format_tour_date(date_birth)

            # Форматируем количество ночей с правильным склонением
            formatted_nights = format_nights(night_num)

            # Определяем значение для лечения (чекбокс "С лечением")
            treatment_text = "с лечением" if with_treatment else "без лечения"

            # Загрузка шаблона
            template_path = resource_path("template.docx")
            doc = Document(template_path)

            # Словарь замен
            replacements = {
                "{full_name}": full_name,
                "{contract_number}": contract_number,
                "{date}": formatted_date,
                "{id_num}": id_num,
                "{tourists_list}": tourists_list,
                "{rooms_list}": rooms_list,
                "{night_num}": formatted_nights,
                "{start_date}": formatted_start_date,
                "{end_date}": formatted_end_date,
                "{date_birth}": formatted_birth_date,
                "{iin}": iin,
                "{treatment}": treatment_text,
                "{treatment_details}": treatment_details_text,
            }

            # Замена плейсхолдеров в параграфах
            for p in doc.paragraphs:
                p.text = replace_placeholders(p.text, replacements)

            # Замена плейсхолдеров в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.text = replace_placeholders(p.text, replacements)

            # Сохранение документа в память
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Отправка файла пользователю
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=f"Договор_{contract_number}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            return f"Ошибка при генерации документа: {str(e)}", 500

    return render_template("form.html")

if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)